"""
ingest_docx.py — Convert DOCX files to Markdown using python-docx.

Part of process-ai ingest pipeline. Uses native heading styles (Heading 1-6)
for reliable structure detection, converts tables to markdown, extracts
embedded images with surrounding paragraph context, and surfaces document
metadata (author, created, revision).

Contract: emits JSON to stdout — { ok, format, pages, markdown, images, metadata }
"""

import argparse
import os
import re

from docx import Document
from docx.oxml.ns import qn

from ingest_common import (
    emit_error,
    emit_success,
    extract_title_from_md,
    slugify_for_path,
    save_image_bytes,
    ensure_images_dir,
)


def _is_heading(paragraph) -> bool:
    """Check if a paragraph has a heading style (Heading 1 through 6)."""
    style_name = (paragraph.style.name if paragraph.style else '')
    if not style_name:
        return False
    return bool(re.match(r'^Heading \d+$', style_name, re.IGNORECASE))


def _heading_level(paragraph) -> int:
    """Extract heading level number from Heading N style."""
    style_name = paragraph.style.name or ''
    m = re.match(r'^Heading (\d+)$', style_name, re.IGNORECASE)
    return int(m.group(1)) if m else 0


def _extract_images(doc, output_dir, slug) -> list:
    """Extract all images from a DOCX file. Returns list of (rel_path, rId, content_type)."""
    images = []
    img_dir = ensure_images_dir(output_dir, slug)
    seen_xrefs = set()

    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        image_part = rel.target_part
        content_type = image_part.content_type  # e.g. 'image/png', 'image/jpeg'
        ext = content_type.split('/')[-1] if '/' in content_type else 'png'
        if ext == 'jpeg':
            ext = 'jpg'

        # Use content hash to dedupe
        data = image_part.blob
        xref = hash(data)
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)

        img_path = save_image_bytes(data, ext, img_dir, len(images))
        images.append((img_path, rel_id))

    return images


def _find_image_rid(paragraph) -> str | None:
    """Find the rId of the first image in a paragraph, if any."""
    for run in paragraph.runs:
        drawings = run._element.findall(qn('w:drawing'))
        for drawing in drawings:
            blip_elements = drawing.findall('.//' + qn('a:blip'))
            for blip in blip_elements:
                rid = blip.get(qn('r:embed'))
                if rid:
                    return rid
    return None


def _table_to_markdown(table) -> str:
    """Convert a docx table to markdown table string."""
    rows = table.rows
    if not rows:
        return ''

    lines = []
    for row_idx, row in enumerate(rows):
        cells = []
        for cell in row.cells:
            text = cell.text.strip().replace('\n', ' ')
            cells.append(text)
        lines.append('| ' + ' | '.join(cells) + ' |')

        if row_idx == 0:
            # Header separator
            lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')

    return '\n'.join(lines)


def convert_docx(docx_path: str, output_dir: str):
    """Convert a DOCX file to Markdown. Returns (md_path, images_rel, metadata)."""
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise RuntimeError(f"Não foi possível abrir o arquivo DOCX (formato inválido?): {e}")

    filename = os.path.basename(docx_path)
    base_name = os.path.splitext(filename)[0]
    slug = slugify_for_path(base_name)

    md_name = f"{slug}.md"
    md_path = os.path.join(output_dir, md_name)

    # Extract images
    img_entries = _extract_images(doc, output_dir, slug)
    img_map = {rid: path for path, rid in img_entries}

    # Extract metadata from core properties
    cp = doc.core_properties
    metadata = {
        'source_file': filename,
        'title': cp.title or base_name,
        'author': cp.author or '',
        'created': cp.created.isoformat() if cp.created else '',
        'modified': cp.modified.isoformat() if cp.modified else '',
    }

    all_images_rel = []
    lines_out = []

    # Process document body in order
    body = doc.element.body
    paragraph_index = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Paragraph
            if paragraph_index >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[paragraph_index]
            paragraph_index += 1

            text = para.text.strip()
            if not text and not _find_image_rid(para):
                continue

            if _is_heading(para):
                level = _heading_level(para)
                lines_out.append('#' * level + ' ' + text + '\n')
            else:
                # Check for inline images
                img_rid = _find_image_rid(para)
                if img_rid and img_rid in img_map:
                    img_rel = os.path.relpath(img_map[img_rid], output_dir).replace('\\', '/')
                    all_images_rel.append(img_rel)
                    has_text = bool(para.text.strip())
                    lines_out.append(f'<figure>')
                    lines_out.append(f'  <img src="{img_rel}" alt="[Image]" />')
                    if has_text:
                        lines_out.append(f'  <figcaption>{para.text.strip()}</figcaption>')
                    else:
                        lines_out.append(f'  <figcaption><!-- [IMAGE] --></figcaption>')
                    lines_out.append(f'</figure>\n')
                elif text:
                    lines_out.append(text + '\n')

        elif tag == 'tbl':
            # Table — find its index among all tables in the document
            tbl_idx = 0
            for prev_child in body:
                if prev_child == child:
                    break
                if prev_child.tag.split('}')[-1] == 'tbl':
                    tbl_idx += 1
            if tbl_idx < len(doc.tables):
                md_table = _table_to_markdown(doc.tables[tbl_idx])
                lines_out.append('\n' + md_table + '\n')

        elif tag == 'sdt':
            # Structured document tag (e.g. TOC) — skip for v1
            pass

    md_body = '\n'.join(lines_out)

    # Cleanup: collapse excessive blank lines
    md_body = re.sub(r'\n{3,}', '\n\n', md_body)

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_body)

    # Extract title from markdown
    extracted_title = extract_title_from_md(md_body)
    if extracted_title:
        metadata['title'] = extracted_title

    # Count "pages" approximately (1 page ≈ 300 words)
    word_count = len(md_body.split())
    est_pages = max(1, word_count // 300) if word_count > 0 else 1

    return md_path, all_images_rel, metadata, est_pages


def main():
    parser = argparse.ArgumentParser(
        description='Convert DOCX to Markdown (process-ai ingest)',
    )
    parser.add_argument('--input', required=True, help='Path to DOCX file')
    parser.add_argument('--output-dir', required=True,
                        help='Directory for output markdown + images')
    args = parser.parse_args()

    docx_path = os.path.abspath(args.input)
    output_dir = os.path.abspath(args.output_dir)

    if not os.path.isfile(docx_path):
        emit_error(f"Arquivo não encontrado: {docx_path}")
        return

    try:
        os.makedirs(output_dir, exist_ok=True)
        md_path, images, metadata, pages = convert_docx(docx_path, output_dir)
        md_rel = os.path.relpath(md_path, output_dir).replace('\\', '/')

        emit_success(
            format='docx',
            markdown_rel=md_rel,
            images=images,
            metadata=metadata,
            pages=pages,
        )
    except Exception as e:
        emit_error(f"Falha ao converter DOCX '{os.path.basename(docx_path)}': {str(e)}")


if __name__ == "__main__":
    main()
